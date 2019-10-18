#! /usr/bin/python3

import docx
import logging
import sys
import readDocx

logging.basicConfig(level=logging.DEBUG, format=' %(asctime)s - %(levelname)s - %(message)s')
#logging.disable(logging.CRITICAL)

guestListFilename = sys.argv[1]
invTempFilename = sys.argv[2]

def getGuestList(guestListFilename):
    with open(guestListFilename) as f:
        guestList = f.readlines()
    return guestList

#def getParagraphs(doc):

def other():
    doc = docx.Document(invTempFilename)
    doc.add_paragraph("Hello World!")
    paraObj1 = doc.add_paragraph("This is a second paragraph")
    paraObj2 = doc.add_paragraph("This is a yet another paragraph")
    paraObj1.add_run("This text is being added to the second paragraph")

    doc.save("helloworld_multipleParagraphs.docx")

if __name__ == "__main__":
    logging.debug(readDocx.getText(invTempFilename))
    gl = getGuestList(guestListFilename)
    logging.debug("Guest List:" )
    logging.debug(gl)

    templateDocList = []
    doc = docx.Document(invTempFilename)

    # Build Document Items
    for i, para in enumerate(doc.paragraphs):
        logging.debug("Index number: " + str(i))
        logging.debug(para.style.name)
        logging.debug(para.text)
        templateDocList.append({"text":para.text,"style":para.style.name,"alignment":para.alignment})
    logging.debug(templateDocList)

    doc2 = docx.Document()
    for guest in gl:
        logging.debug("Working on invitation for: " + guest)
        for paragraph in templateDocList:
            if paragraph['text'] == "[NAME]":
                doc2.add_paragraph(guest,style=paragraph['style'])
            else:
                doc2.add_paragraph(paragraph['text'],style=paragraph['style'])
        doc2.add_page_break()
    doc2.save('FinishedInvitations.docx')


